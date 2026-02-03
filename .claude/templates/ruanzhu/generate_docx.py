#!/usr/bin/env python3
"""
软著源代码DOCX生成器

支持语言：Java, TypeScript/Vue, C++, Ruby, Rust, Go, Python

用法：
    python generate_docx.py [选项]

选项：
    --name, -n     软件名称（默认自动检测）
    --version, -v  版本号（默认 V1.0）
    --pages, -p    目标页数或 auto（默认 60）
    --root, -r     项目根目录（默认当前目录）

示例：
    python generate_docx.py
    python generate_docx.py --name "智能仓储系统" --version "V2.0"
    python generate_docx.py --pages auto

输出：docs/ruanzhu/{软件名称}{版本}-源代码.docx
依赖：python-docx（会自动创建虚拟环境安装）
"""

import os
import re
import sys
import math
import json
import fnmatch
import argparse
import subprocess
from pathlib import Path

# ============== 配置 ==============

# 语言配置：扩展名 -> (语言名, 优先目录列表, 排除模式列表)
LANGUAGE_CONFIG = {
    'java': {
        'extensions': ['.java'],
        'priority_dirs': [
            'src/main/java/**/controller',
            'src/main/java/**/service',
            'src/main/java/**/entity',
            'src/main/java/**/repository',
            'src/main/java/**/config',
            'src/main/java/**/security',
            'src/main/java/**/dto',
            'src/main/java',
        ],
        'exclude_patterns': ['*Test.java', '*IT.java', '*Tests.java'],
        'exclude_dirs': ['target', 'build', '.gradle'],
    },
    'typescript': {
        'extensions': ['.ts', '.tsx', '.vue'],
        'priority_dirs': [
            'src/api',
            'src/stores',
            'src/pages',
            'src/views',
            'src/components',
            'src/hooks',
            'src/utils',
            'src/layouts',
            'frontend/src/api',
            'frontend/src/stores',
            'frontend/src/pages',
            'frontend/src/views',
            'frontend/src/components',
            'frontend/src',
            'src',
        ],
        'exclude_patterns': ['*.spec.ts', '*.test.ts', '*.spec.tsx', '*.test.tsx', '*.d.ts'],
        'exclude_dirs': ['node_modules', 'dist', 'build', '.next'],
    },
    'cpp': {
        'extensions': ['.cpp', '.cc', '.cxx', '.hpp', '.h'],
        'priority_dirs': ['src', 'include', 'lib'],
        'exclude_patterns': ['*_test.cpp', '*_test.cc', '*_test.h'],
        'exclude_dirs': ['build', 'cmake-build-debug', 'cmake-build-release', 'test', 'tests'],
    },
    'ruby': {
        'extensions': ['.rb'],
        'priority_dirs': ['app/controllers', 'app/models', 'app/services', 'lib', 'app'],
        'exclude_patterns': ['*_spec.rb', '*_test.rb'],
        'exclude_dirs': ['spec', 'test', 'vendor'],
    },
    'rust': {
        'extensions': ['.rs'],
        'priority_dirs': ['src'],
        'exclude_patterns': ['*_test.rs'],
        'exclude_dirs': ['target', 'tests'],
    },
    'go': {
        'extensions': ['.go'],
        'priority_dirs': ['cmd', 'internal', 'pkg', '.'],
        'exclude_patterns': ['*_test.go'],
        'exclude_dirs': ['vendor', 'testdata'],
    },
    'python': {
        'extensions': ['.py'],
        'priority_dirs': ['src', 'app', 'lib', '.'],
        'exclude_patterns': ['test_*.py', '*_test.py', 'conftest.py'],
        'exclude_dirs': ['tests', 'test', '__pycache__', '.venv', 'venv', '.pytest_cache'],
    },
}

# 项目检测文件
PROJECT_DETECT = {
    'pom.xml': 'java',
    'build.gradle': 'java',
    'package.json': 'typescript',
    'Cargo.toml': 'rust',
    'Gemfile': 'ruby',
    'go.mod': 'go',
    'CMakeLists.txt': 'cpp',
    'requirements.txt': 'python',
    'pyproject.toml': 'python',
}

# DOCX格式配置
DOCX_CONFIG = {
    'page_width_cm': 21,
    'page_height_cm': 29.7,
    'margin_top_cm': 2.5,
    'margin_bottom_cm': 2.5,
    'margin_left_cm': 3.0,
    'margin_right_cm': 2.5,
    'font_size_pt': 10,
    'font_cn': '宋体',
    'font_en': 'Courier New',
    'lines_per_page': 57,
}


def ensure_docx_lib():
    """确保python-docx库可用"""
    try:
        import docx
        return True
    except ImportError:
        print("python-docx 未安装，正在创建虚拟环境...")
        venv_path = Path(__file__).parent / '.venv'
        if not venv_path.exists():
            subprocess.run([sys.executable, '-m', 'venv', str(venv_path)], check=True)

        pip_path = venv_path / 'bin' / 'pip'
        if not pip_path.exists():
            pip_path = venv_path / 'Scripts' / 'pip.exe'

        subprocess.run([str(pip_path), 'install', 'python-docx', '-q'], check=True)

        # 重新执行脚本
        python_path = venv_path / 'bin' / 'python'
        if not python_path.exists():
            python_path = venv_path / 'Scripts' / 'python.exe'
        os.execv(str(python_path), [str(python_path)] + sys.argv)


def detect_project_languages(project_root):
    """检测项目使用的语言"""
    detected = set()
    root = Path(project_root)

    # 检查根目录
    for filename, lang in PROJECT_DETECT.items():
        if (root / filename).exists():
            detected.add(lang)

    # 检查常见子目录（前端项目常在子目录）
    for subdir in ['frontend', 'client', 'web', 'app', 'ui']:
        subpath = root / subdir
        if subpath.exists():
            for filename, lang in PROJECT_DETECT.items():
                if (subpath / filename).exists():
                    detected.add(lang)

    # 如果没检测到，扫描文件扩展名
    if not detected:
        for ext_lang, config in LANGUAGE_CONFIG.items():
            for ext in config['extensions']:
                if list(root.rglob(f'*{ext}'))[:1]:
                    detected.add(ext_lang)
                    break

    return list(detected)


def read_project_info(project_root):
    """读取项目名称和版本"""
    root = Path(project_root)
    name, version = None, None

    # 1. 尝试从 CLAUDE.md 读取
    for claude_path in [root / '.claude' / 'CLAUDE.md', root / 'CLAUDE.md']:
        if claude_path.exists():
            content = claude_path.read_text(encoding='utf-8')
            # 匹配 "# 项目名" 或 "## 项目概述" 后的内容
            name_match = re.search(r'^#\s+(.+?)(?:\s*[-–—]|\n)', content, re.MULTILINE)
            if name_match:
                name = name_match.group(1).strip()
            version_match = re.search(r'版本[：:]\s*[vV]?([\d.]+)', content)
            if version_match:
                version = f"V{version_match.group(1)}"
            if name:
                break

    # 2. 尝试从 package.json 读取
    if not name:
        pkg_path = root / 'package.json'
        if pkg_path.exists():
            try:
                pkg = json.loads(pkg_path.read_text(encoding='utf-8'))
                name = pkg.get('name', '').replace('-', ' ').replace('_', ' ').title()
                version = f"V{pkg.get('version', '1.0')}"
            except (json.JSONDecodeError, OSError):
                pass

    # 3. 尝试从 pom.xml 读取
    if not name:
        pom_path = root / 'pom.xml'
        if pom_path.exists():
            try:
                content = pom_path.read_text(encoding='utf-8')
                art_match = re.search(r'<artifactId>([^<]+)</artifactId>', content)
                ver_match = re.search(r'<version>([^<]+)</version>', content)
                if art_match:
                    name = art_match.group(1).replace('-', ' ').replace('_', ' ').title()
                if ver_match:
                    version = f"V{ver_match.group(1)}"
            except OSError:
                pass

    return name, version or 'V1.0'


def should_exclude(filepath, config):
    """检查文件是否应该排除"""
    path = Path(filepath)

    # 检查目录排除
    for exclude_dir in config.get('exclude_dirs', []):
        if exclude_dir in path.parts:
            return True

    # 检查文件名模式排除（使用 fnmatch）
    filename = path.name
    for pattern in config.get('exclude_patterns', []):
        if fnmatch.fnmatch(filename, pattern):
            return True

    return False


def collect_source_files(project_root, languages):
    """收集源代码文件"""
    root = Path(project_root)
    files = []

    for lang in languages:
        config = LANGUAGE_CONFIG.get(lang, {})
        extensions = config.get('extensions', [])
        priority_dirs = config.get('priority_dirs', ['.'])

        # 按优先目录顺序收集
        for pdir in priority_dirs:
            if '**' in pdir:
                # 通配符目录
                base, pattern = pdir.split('**')
                base_path = root / base.rstrip('/')
                if base_path.exists():
                    for subdir in base_path.iterdir():
                        if subdir.is_dir() and subdir.name.endswith(pattern.lstrip('/')):
                            for ext in extensions:
                                for f in sorted(subdir.rglob(f'*{ext}')):
                                    if not should_exclude(f, config):
                                        files.append((f.name, lang, f))
            else:
                dir_path = root / pdir
                if dir_path.exists() and dir_path.is_dir():
                    for ext in extensions:
                        for f in sorted(dir_path.rglob(f'*{ext}')):
                            if not should_exclude(f, config):
                                files.append((f.name, lang, f))

    # 去重
    seen = set()
    unique_files = []
    for name, lang, path in files:
        if path not in seen:
            seen.add(path)
            unique_files.append((name, lang, path))

    return unique_files


def read_file_lines(filepath):
    """读取文件内容"""
    try:
        content = Path(filepath).read_text(encoding='utf-8')
        lines = content.split('\n')
        while lines and lines[-1].strip() == '':
            lines.pop()
        return lines
    except (OSError, UnicodeDecodeError):
        return []


def estimate_rendered_lines(text, usable_width_cm=15.5):
    """估算一行文本在Word中渲染需要的行数"""
    if len(text) == 0:
        return 1
    width = 0
    for ch in text:
        if ord(ch) > 127:
            width += 0.35
        else:
            width += 0.18
    return max(1, int(width / usable_width_cm) + (1 if width % usable_width_cm > 0.1 else 0))


def build_source_lines(source_files, target_rendered_lines, mode='fixed'):
    """
    构建源代码行

    Args:
        source_files: 源文件列表 [(filename, lang, path), ...]
        target_rendered_lines: 目标渲染行数
        mode: 'fixed' 固定页数, 'auto' 自动模式
    """
    all_lines = []
    total_rendered = 0

    for filename, _lang, filepath in source_files:
        lines = read_file_lines(filepath)
        if not lines:
            continue

        # 文件标识注释
        header_line = f"/* ========== {filename} ========== */"
        header_rendered = estimate_rendered_lines(header_line)

        if mode == 'fixed' and total_rendered + header_rendered + 2 > target_rendered_lines:
            break

        all_lines.append(header_line)
        all_lines.append("")
        total_rendered += header_rendered + 1

        for line in lines:
            line_rendered = estimate_rendered_lines(line)
            if mode == 'fixed' and total_rendered + line_rendered > target_rendered_lines:
                break
            all_lines.append(line)
            total_rendered += line_rendered

        all_lines.append("")
        total_rendered += 1

        if mode == 'fixed' and total_rendered >= target_rendered_lines:
            break

    return all_lines, total_rendered


def generate_docx(lines, output_path, software_name, version, is_split=False, total_pages=0):
    """生成DOCX文件"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    cfg = DOCX_CONFIG

    # 设置页面
    section = doc.sections[0]
    section.page_width = Cm(cfg['page_width_cm'])
    section.page_height = Cm(cfg['page_height_cm'])
    section.top_margin = Cm(cfg['margin_top_cm'])
    section.bottom_margin = Cm(cfg['margin_bottom_cm'])
    section.left_margin = Cm(cfg['margin_left_cm'])
    section.right_margin = Cm(cfg['margin_right_cm'])

    # 设置页眉
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]

    # 页眉段落属性
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '8640')
    tabs.append(tab)
    pPr.append(tabs)

    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 左侧标题
    run1 = p.add_run(f"{software_name} {version}")
    run1.font.name = cfg['font_en']
    run1.font.size = Pt(cfg['font_size_pt'])
    run1.font.bold = True
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

    # Tab
    run_tab = p.add_run()
    tab_elem = OxmlElement('w:tab')
    run_tab._element.append(tab_elem)

    # 右侧页码
    def add_field(para, field_code):
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        run2 = para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f" {field_code} "
        run2._element.append(instrText)
        run3 = para.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

    run2 = p.add_run("第 ")
    run2.font.name = cfg['font_en']
    run2.font.size = Pt(cfg['font_size_pt'])
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    add_field(p, "PAGE")
    run3 = p.add_run(" 页共 ")
    run3.font.name = cfg['font_en']
    run3.font.size = Pt(cfg['font_size_pt'])
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    add_field(p, "NUMPAGES")
    run4 = p.add_run(" 页")
    run4.font.name = cfg['font_en']
    run4.font.size = Pt(cfg['font_size_pt'])
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

    # 设置页脚
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run("第 ")
    run1.font.name = cfg['font_en']
    run1.font.size = Pt(cfg['font_size_pt'])
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    add_field(p, "PAGE")
    run2 = p.add_run(" 页")
    run2.font.name = cfg['font_en']
    run2.font.size = Pt(cfg['font_size_pt'])
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.size = Pt(cfg['font_size_pt'])
    style.font.name = cfg['font_en']
    style.element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # 添加代码行
    if is_split:
        front_lines, back_lines = lines
        for line in front_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(line)
            run.font.size = Pt(cfg['font_size_pt'])
            run.font.name = cfg['font_en']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

        # 分隔标记
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n（以上为源代码前30页，以下为源代码后30页，中间省略部分共{total_pages - 60}页）\n")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = cfg['font_en']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

        for line in back_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(line)
            run.font.size = Pt(cfg['font_size_pt'])
            run.font.name = cfg['font_en']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])
    else:
        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(line)
            run.font.size = Pt(cfg['font_size_pt'])
            run.font.name = cfg['font_en']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), cfg['font_cn'])

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='软著源代码DOCX生成器')
    parser.add_argument('--name', '-n', help='软件名称')
    parser.add_argument('--version', '-v', default='V1.0', help='版本号 (默认: V1.0)')
    parser.add_argument('--pages', '-p', default='60', help='目标页数或auto (默认: 60)')
    parser.add_argument('--root', '-r', default='.', help='项目根目录 (默认: 当前目录)')
    args = parser.parse_args()

    # 确保依赖
    ensure_docx_lib()

    project_root = Path(args.root).resolve()
    print(f"项目目录: {project_root}")

    # 检测项目信息
    detected_name, detected_version = read_project_info(project_root)
    software_name = args.name or detected_name
    version = args.version if args.version != 'V1.0' else detected_version

    if not software_name:
        software_name = input("请输入软件名称: ").strip()
        if not software_name:
            print("错误: 未指定软件名称")
            sys.exit(1)

    print(f"软件名称: {software_name} {version}")

    # 检测语言
    languages = detect_project_languages(project_root)
    if not languages:
        print("错误: 未检测到支持的编程语言")
        print(f"支持的语言: {', '.join(LANGUAGE_CONFIG.keys())}")
        sys.exit(1)
    print(f"检测到语言: {', '.join(languages)}")

    # 收集源文件
    print("扫描源代码文件...")
    source_files = collect_source_files(project_root, languages)
    print(f"找到 {len(source_files)} 个源文件")

    if not source_files:
        print("错误: 未找到源代码文件")
        sys.exit(1)

    # 计算目标行数
    lines_per_page = DOCX_CONFIG['lines_per_page']

    if args.pages.lower() == 'auto':
        # 自动模式：先收集全部，再决定是否分割
        all_lines, total_rendered = build_source_lines(source_files, float('inf'), mode='auto')
        total_pages = math.ceil(total_rendered / lines_per_page)
        print(f"总渲染行数: {total_rendered}, 总页数: {total_pages}")

        if total_pages <= 60:
            lines = all_lines
            is_split = False
        else:
            # 分割：前30页 + 后30页
            front_lines = lines_per_page * 30
            back_lines = lines_per_page * 30

            front = []
            front_count = 0
            for line in all_lines:
                rendered = estimate_rendered_lines(line)
                if front_count + rendered > front_lines:
                    break
                front.append(line)
                front_count += rendered

            back = []
            back_count = 0
            for line in reversed(all_lines):
                rendered = estimate_rendered_lines(line)
                if back_count + rendered > back_lines:
                    break
                back.insert(0, line)
                back_count += rendered

            lines = (front, back)
            is_split = True
            print(f"自动模式: 前30页 + 后30页 (省略 {total_pages - 60} 页)")
    else:
        try:
            target_pages = int(args.pages)
        except ValueError:
            print(f"错误: 页数参数无效: {args.pages}")
            sys.exit(1)
        target_rendered = target_pages * lines_per_page
        lines, total_rendered = build_source_lines(source_files, target_rendered, mode='fixed')
        actual_pages = math.ceil(total_rendered / lines_per_page)
        print(f"目标 {target_pages} 页, 实际约 {actual_pages} 页 ({total_rendered} 渲染行)")
        is_split = False
        total_pages = actual_pages

    # 创建输出目录
    output_dir = project_root / 'docs' / 'ruanzhu'
    output_dir.mkdir(parents=True, exist_ok=True)

    # 生成文件名
    safe_name = re.sub(r'[<>:"/\\|?*]', '', software_name)
    output_path = output_dir / f"{safe_name}{version}-源代码.docx"

    print(f"生成DOCX: {output_path}")
    generate_docx(lines, str(output_path), software_name, version, is_split, total_pages if is_split else 0)

    print(f"\n完成! 文件: {output_path}")
    print(f"页眉: {software_name} {version}")
    print(f"请在Word中打开确认页数")


if __name__ == '__main__':
    main()
